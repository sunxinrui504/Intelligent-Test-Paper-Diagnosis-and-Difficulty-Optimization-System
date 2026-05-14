# exporter.py
import os
from docx import Document
from docx.shared import Pt, RGBColor
import pandas as pd

class ReportExporter:
    @staticmethod
    def export_exam_to_word(exam_text: str, filename: str, output_dir: str):
        """將 AI 生成的試卷匯出為格式化的 Word (.docx) 文件"""
        doc = Document()
        doc.add_heading(f"AI Generated Exam - {filename.split('_')[1]} Version", 0)
        
        for paragraph in exam_text.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                # 簡單的樣式處理
                if paragraph.startswith('一、') or paragraph.startswith('二、'):
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(14)
                
                # 模擬 PaddleDetection/圖片佔位符的處理
                if "[圖片佔位符" in paragraph:
                    p.runs[0].font.color.rgb = RGBColor(0xFF, 0, 0) # 標紅顯示
                    
        save_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(save_path)
        return save_path

    @staticmethod
    def generate_echarts_html(hist_df, best_model_name, target_dcw, output_dir):
        """生成真實的 ECharts HTML 檔案"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>AI Exam Analytics ECharts Dashboard</title>
            <script src="https://cdn.jsdelivr.net/npm/echarts@5.4.3/dist/echarts.min.js"></script>
        </head>
        <body>
            <h2 style="text-align:center;">Enterprise Exam Analytics Dashboard</h2>
            <p style="text-align:center;">Winning ML Model: <b>{best_model_name}</b> | Target DCW: <b>{target_dcw:.2f}</b></p>
            <div id="scatterChart" style="width: 800px;height:400px; margin:auto;"></div>
            <script>
                var chartDom = document.getElementById('scatterChart');
                var myChart = echarts.init(chartDom);
                var option = {{
                    title: {{ text: 'Difficulty vs Score Prediction' }},
                    tooltip: {{ trigger: 'axis' }},
                    xAxis: {{ name: 'DCW (Difficulty)', type: 'value' }},
                    yAxis: {{ name: 'Average Score', type: 'value' }},
                    series: [{{
                        symbolSize: 10,
                        data: {hist_df[['DCW', 'Average_Score']].values.tolist()},
                        type: 'scatter'
                    }}]
                }};
                myChart.setOption(option);
            </script>
        </body>
        </html>
        """
        path = os.path.join(output_dir, "ECharts_Dashboard.html")
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return path
